import os
import tempfile
from io import BytesIO
from docx import Document
import pandas as pd
import datetime
import webvtt
import streamlit as st


def vtt_to_dataframe(file):
    """
    Convert a VTT file to a DataFrame.

    Args:
        file: A VTT file-like object.

    Returns:
        DataFrame containing the captions from the VTT file.
    """
    # Read the content of the uploaded file
    content = file.read().decode("utf-8")

    # Create a temporary file to store the content
    with tempfile.NamedTemporaryFile(delete=False, suffix=".vtt") as temp_file:
        temp_file.write(content.encode("utf-8"))
        temp_file_path = temp_file.name  # Store the temporary file path

    # Use webvtt to read the temporary file
    rows = []
    for caption in webvtt.read(temp_file_path):
        rows.append({
            'Start': caption.start,
            'End': caption.end,
            'Speaker': caption.voice,
            'Text': caption.text
        })

    return pd.DataFrame(rows)


def map_speakers(dataframe):
    """
    Map speakers to numerical labels.

    Args:
        dataframe: A DataFrame containing captions.

    Returns:
        DataFrame with a new 'Label' column for speakers.
    """
    speakers_dict = {speaker: index for index, speaker in enumerate(dataframe["Speaker"].unique())}
    dataframe["Label"] = dataframe["Speaker"].map(speakers_dict)
    return dataframe


def remove_breaks_text(dataframe):
    """
    Remove line breaks from the text.

    Args:
        dataframe: A DataFrame containing captions.

    Returns:
        DataFrame with cleaned text.
    """
    dataframe['Text'] = dataframe['Text'].apply(lambda x: str(x).replace('\n', ' '))
    return dataframe


def concatenate_text_with_timestamp_and_speaker_by_label(dataframe):
    """
    Concatenate text based on timestamps and speakers.

    Args:
        dataframe: A DataFrame containing captions.

    Returns:
        DataFrame with concatenated text entries.
    """
    concatenated_data = []
    current_text = ""
    start_timestamp = None
    speaker = None

    for index, row in dataframe.iterrows():
        if start_timestamp is None or row["Label"] == dataframe.at[index - 1, "Label"]:
            if start_timestamp is None:
                start_timestamp = row["Start"]
            if speaker is None:
                speaker = row["Speaker"]
            current_text += " " + row["Text"]
            end_timestamp = row["End"]
        else:
            concatenated_data.append({
                "text": current_text.strip(),
                "start_timestamp": start_timestamp,
                "end_timestamp": end_timestamp,
                "speaker": speaker,
            })
            current_text = row["Text"]
            start_timestamp = row["Start"]
            end_timestamp = row["End"]
            speaker = row["Speaker"]

    # Add the last concatenated entry
    concatenated_data.append({
        "text": current_text.strip(),
        "start_timestamp": start_timestamp,
        "end_timestamp": end_timestamp,
        "speaker": speaker,
    })
    return pd.DataFrame(concatenated_data)


def write_to_word_doc(concatenated_data):
    """
    Write the concatenated text data to a Word document.

    Args:
        concatenated_data: A DataFrame with concatenated captions.

    Returns:
        A Document object containing the formatted text.
    """
    document = Document()
    document.add_heading("Meeting Transcription", level=0)

    for index, row in concatenated_data.iterrows():
        document.add_paragraph(f"[{row['start_timestamp']} -- {row['end_timestamp']}]")
        speaker_paragraph = document.add_paragraph()
        speaker_paragraph.add_run(f"Speaker: {row['speaker']}").bold = True
        document.add_paragraph(row["text"])

    return document


def main():
    """
    Main function to run the Streamlit app.
    """
    st.title("VTT to Word Document Converter")

    # Allow multiple file uploads
    uploaded_files = st.file_uploader("Choose VTT files", type=["vtt"], accept_multiple_files=True)

    if uploaded_files:
        # Process each uploaded VTT file
        for uploaded_file in uploaded_files:
            # Read the file into a DataFrame
            df = vtt_to_dataframe(uploaded_file)
            df = map_speakers(df)
            df = remove_breaks_text(df)
            df = concatenate_text_with_timestamp_and_speaker_by_label(df)

            # Write the document to a BytesIO stream
            doc = write_to_word_doc(df)
            output_file_name = f"{os.path.splitext(uploaded_file.name)[0]}-FORMATTED-{datetime.datetime.now().strftime('%Y-%m-%d')}.docx"

            # Save to BytesIO
            doc_stream = BytesIO()
            doc.save(doc_stream)
            doc_stream.seek(0)  # Move to the beginning of the stream

            # Provide a download button for each file
            st.download_button(
                label=f"Download {output_file_name}",
                data=doc_stream,
                file_name=output_file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )


if __name__ == '__main__':
    main()
